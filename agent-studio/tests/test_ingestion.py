import base64
import io
from types import SimpleNamespace

import pytest
from fastapi.testclient import TestClient

import agent_studio.main as main_module
from agent_studio.config import get_settings
from agent_studio.db import TrustedContext
from agent_studio.ingestion import ExtractionError, extract_file
from agent_studio.integration_config import integration_config_store
from agent_studio.main import app, knowledge_ingestion_service, knowledge_ingestion_store
from agent_studio.schemas import KnowledgeHit, KnowledgeIngestionFile
from agent_studio.store import store


client = TestClient(app)


def setup_function() -> None:
    store.clear()
    integration_config_store.clear()
    knowledge_ingestion_store.clear()
    get_settings.cache_clear()


def _base64_bytes(raw: bytes) -> str:
    return base64.b64encode(raw).decode("ascii")


def _search_excerpts(query: str) -> list[str]:
    response = client.post(
        "/knowledge/search-test",
        json={"query": query, "intent": "general_support", "risk_level": "medium"},
    )
    assert response.status_code == 200
    return [hit["excerpt"].lower() for hit in response.json()["hits"]]


def test_extractors_parse_supported_fixtures() -> None:
    markdown = extract_file(
        KnowledgeIngestionFile(filename="refund-policy.md", content="# Refund Policy\nSale items need review."),
    )
    assert markdown.title == "Refund Policy"
    assert "Sale items" in markdown.content

    transcript = extract_file(
        KnowledgeIngestionFile(
            filename="call-transcript.json",
            content='{"messages":[{"speaker":"Customer","text":"My order is late."}]}',
        ),
    )
    assert "Customer: My order is late." in transcript.content

    csv_doc = extract_file(
        KnowledgeIngestionFile(
            filename="shipping-faq.csv",
            content="Question,Answer\nWhere is my order?,Check tracking first.",
        ),
    )
    assert "Question: Where is my order?" in csv_doc.content

    docx_buffer = io.BytesIO()
    from docx import Document

    docx = Document()
    docx.add_heading("Warranty SOP", level=1)
    docx.add_paragraph("Ask for proof of purchase before warranty action.")
    docx.save(docx_buffer)
    docx_doc = extract_file(
        KnowledgeIngestionFile(
            filename="warranty-sop.docx",
            content=_base64_bytes(docx_buffer.getvalue()),
            encoding="base64",
        ),
    )
    assert "proof of purchase" in docx_doc.content

    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Returns"
    sheet.append(["Scenario", "Rule"])
    sheet.append(["Sale item", "Escalate before promising refund"])
    xlsx_buffer = io.BytesIO()
    workbook.save(xlsx_buffer)
    xlsx_doc = extract_file(
        KnowledgeIngestionFile(
            filename="returns.xlsx",
            content=_base64_bytes(xlsx_buffer.getvalue()),
            encoding="base64",
        ),
    )
    assert "Scenario: Sale item" in xlsx_doc.content

    pdf_doc = extract_file(
        KnowledgeIngestionFile(
            filename="policy.pdf",
            content=_base64_bytes(b"%PDF-1.4\nBT (PDF refund policy text) Tj ET\n%%EOF"),
            encoding="base64",
        ),
    )
    assert "PDF refund policy text" in pdf_doc.content


def test_scanned_pdf_without_text_reports_ocr_needed(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("SAGAD_OCR_ENABLED", "false")
    monkeypatch.setenv("SAGAD_DOCLING_ENABLED", "false")
    get_settings.cache_clear()

    with pytest.raises(ExtractionError) as exc_info:
        extract_file(
            KnowledgeIngestionFile(
                filename="scan.pdf",
                content=_base64_bytes(b"%PDF-1.4\n%%EOF"),
                encoding="base64",
            ),
        )

    assert exc_info.value.code == "ocr_required"
    assert "OCR" in exc_info.value.message



def test_scanned_pdf_uses_ocr_when_enabled(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("SAGAD_OCR_ENABLED", "true")
    monkeypatch.setenv("SAGAD_OCR_LANG", "eng")
    get_settings.cache_clear()

    monkeypatch.setattr(
        "agent_studio.ingestion._convert_pdf_to_images",
        lambda raw, max_pages: [SimpleNamespace(page_number=1)],
        raising=False,
    )
    monkeypatch.setattr(
        "agent_studio.ingestion._image_to_text",
        lambda image, lang, timeout_seconds: "Scanned refund SOP text",
        raising=False,
    )

    document = extract_file(
        KnowledgeIngestionFile(
            filename="scan.pdf",
            content=_base64_bytes(b"%PDF-1.4\n%%EOF"),
            encoding="base64",
        ),
    )

    assert "Scanned refund SOP text" in document.content
    assert document.metadata["ocr_used"] is True
    assert document.metadata["ocr_pages"] == 1
    assert document.metadata["ocr_lang"] == "eng"


def test_scanned_pdf_reports_ocr_unavailable_when_enabled_but_runtime_missing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("SAGAD_OCR_ENABLED", "true")
    get_settings.cache_clear()

    def fail_conversion(raw: bytes, max_pages: int) -> list[object]:
        raise RuntimeError("poppler is not installed")

    monkeypatch.setattr(
        "agent_studio.ingestion._convert_pdf_to_images",
        fail_conversion,
        raising=False,
    )

    with pytest.raises(ExtractionError) as exc_info:
        extract_file(
            KnowledgeIngestionFile(
                filename="scan.pdf",
                content=_base64_bytes(b"%PDF-1.4\n%%EOF"),
                encoding="base64",
            ),
        )

    assert exc_info.value.code == "ocr_unavailable"
    assert "poppler" in exc_info.value.message


def test_ingestion_job_creates_needs_review_document() -> None:
    response = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Test Upload",
            "files": [
                {
                    "filename": "shipping-faq.md",
                    "content": "# Shipping FAQ\nUse tracking before escalation.",
                },
            ],
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["job"]["status"] == "needs_review"
    assert payload["job"]["processed_files"] == 1
    assert payload["documents"][0]["approval_status"] == "needs_review"
    assert payload["documents"][0]["chunk_count"] == 0


def test_unapproved_document_does_not_retrieve_until_approved() -> None:
    created = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Support Upload",
            "files": [
                {
                    "filename": "zephyrwidget-sop.md",
                    "category": "sops",
                    "content": "# ZephyrWidget SOP\nUse zephyrwidget-token before approving service credits.",
                },
            ],
        },
    ).json()
    document_id = created["documents"][0]["id"]

    before = client.post(
        "/knowledge/search-test",
        json={"query": "zephyrwidget-token", "intent": "general_support", "risk_level": "medium"},
    ).json()
    assert all("zephyrwidget" not in hit["excerpt"].lower() for hit in before["hits"])

    approved = client.post(f"/knowledge/documents/{document_id}/approve")
    assert approved.status_code == 200
    assert approved.json()["approval_status"] == "approved"
    assert approved.json()["chunk_count"] >= 1

    after = client.post(
        "/knowledge/search-test",
        json={"query": "zephyrwidget-token", "intent": "general_support", "risk_level": "medium"},
    ).json()
    assert any("zephyrwidget" in hit["excerpt"].lower() for hit in after["hits"])


def test_duplicate_upload_does_not_create_duplicate_active_document() -> None:
    payload = {
        "source_name": "Duplicate Upload",
        "files": [
            {
                "filename": "refund-policy.md",
                "source_path": "policies/refund-policy.md",
                "content": "# Refund Policy\nDuplicate uploads should collapse.",
            },
            {
                "filename": "refund-policy-copy.md",
                "source_path": "policies/refund-policy.md",
                "content": "# Refund Policy\nDuplicate uploads should collapse.",
            },
        ],
    }

    response = client.post("/knowledge/ingestion-jobs", json=payload)
    assert response.status_code == 200

    documents = client.get("/knowledge/documents").json()["documents"]
    matching = [
        document
        for document in documents
        if document["source_path"] == "policies/refund-policy.md"
    ]
    assert len(matching) == 1


def test_same_path_from_different_sources_keeps_distinct_documents() -> None:
    for source_name in ("Client A Upload", "Client B Upload"):
        response = client.post(
            "/knowledge/ingestion-jobs",
            json={
                "source_name": source_name,
                "files": [
                    {
                        "filename": "shared.md",
                        "source_path": "policies/shared.md",
                        "content": f"# Shared SOP\nUnique content for {source_name}.",
                    },
                ],
            },
        )
        assert response.status_code == 200

    documents = client.get("/knowledge/documents").json()["documents"]
    matching = [
        document
        for document in documents
        if document["source_path"] == "policies/shared.md"
    ]

    assert len(matching) == 2
    assert {document["source_id"] for document in matching}


def test_sources_endpoint_lists_manual_ingestion_source() -> None:
    created = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Source Upload",
            "files": [
                {
                    "filename": "source-faq.md",
                    "content": "# Source FAQ\nUse source-specific answer.",
                },
            ],
        },
    ).json()

    response = client.get("/knowledge/sources")

    assert response.status_code == 200
    sources = response.json()["sources"]
    assert any(source["id"] == created["job"]["source_id"] for source in sources)
    assert any(source["name"] == "Source Upload" for source in sources)


def test_document_resync_refreshes_approved_chunks() -> None:
    created = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Resync Upload",
            "files": [
                {
                    "filename": "resync-sop.md",
                    "content": "# Resync SOP\nRefresh this approved answer source.",
                },
            ],
        },
    ).json()
    document_id = created["documents"][0]["id"]
    approved = client.post(f"/knowledge/documents/{document_id}/approve").json()
    assert approved["approval_status"] == "approved"
    assert approved["chunk_count"] >= 1

    response = client.post(f"/knowledge/documents/{document_id}/resync")

    assert response.status_code == 200
    payload = response.json()
    assert payload["approval_status"] == "approved"
    assert payload["chunk_count"] >= 1
    assert "last_resynced_at" in payload["metadata"]


def test_source_sync_refreshes_documents_for_source() -> None:
    created = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Sync Upload",
            "files": [
                {
                    "filename": "sync-faq.md",
                    "content": "# Sync FAQ\nRefresh the local source.",
                },
            ],
        },
    ).json()
    source_id = created["job"]["source_id"]

    response = client.post(f"/knowledge/sources/{source_id}/sync")

    assert response.status_code == 200
    payload = response.json()
    assert payload["job"]["source_id"] == source_id
    assert payload["job"]["processed_files"] == 1
    assert payload["documents"][0]["metadata"]["last_resynced_at"]


def test_search_test_forwards_trusted_context_to_retriever(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured_context: TrustedContext | None = None

    class SpyRetriever:
        records: list[object] = []

        def add_record(self, record: object) -> None:
            pass

        def remove_record(self, record_id: str) -> None:
            pass

        def search(
            self,
            query: str,
            *,
            intent: str,
            risk_level: str,
            limit: int = 4,
            context: TrustedContext | None = None,
        ) -> list[KnowledgeHit]:
            nonlocal captured_context
            captured_context = context
            return []

    monkeypatch.setattr(main_module, "retriever", SpyRetriever())

    response = client.post(
        "/knowledge/search-test",
        headers={
            "x-sagad-org-id": "org-custom",
            "x-sagad-user-id": "42",
            "x-sagad-role": "qa_analyst",
        },
        json={
            "query": "refund sale items",
            "intent": "general_support",
            "risk_level": "medium",
        },
    )

    assert response.status_code == 200
    assert captured_context == TrustedContext(
        organization_id="org-custom",
        user_id="42",
        role="qa_analyst",
    )


def test_archive_removes_approved_document_from_runtime_retrieval() -> None:
    created = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Archive Upload",
            "files": [
                {
                    "filename": "archive-sop.md",
                    "content": "# Archive SOP\nUse archivecontext-token for this retired policy.",
                },
            ],
        },
    ).json()
    document_id = created["documents"][0]["id"]

    approved = client.post(f"/knowledge/documents/{document_id}/approve")
    assert approved.status_code == 200
    assert any("archivecontext-token" in excerpt for excerpt in _search_excerpts("archivecontext-token"))

    archived = client.post(f"/knowledge/documents/{document_id}/archive")
    assert archived.status_code == 200
    assert archived.json()["approval_status"] == "archived"

    assert all("archivecontext-token" not in excerpt for excerpt in _search_excerpts("archivecontext-token"))


def test_changed_reupload_removes_stale_approved_runtime_record() -> None:
    first = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Changed Upload",
            "files": [
                {
                    "filename": "changed-sop.md",
                    "source_path": "policies/changed-sop.md",
                    "content": "# Changed SOP\nUse oldcontext-token for the old policy.",
                },
            ],
        },
    ).json()
    document_id = first["documents"][0]["id"]

    approved = client.post(f"/knowledge/documents/{document_id}/approve")
    assert approved.status_code == 200
    assert any("oldcontext-token" in excerpt for excerpt in _search_excerpts("oldcontext-token"))

    second = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Changed Upload",
            "files": [
                {
                    "filename": "changed-sop.md",
                    "source_path": "policies/changed-sop.md",
                    "content": "# Changed SOP\nUse newcontext-token only after review.",
                },
            ],
        },
    )
    assert second.status_code == 200
    assert second.json()["documents"][0]["id"] == document_id
    assert second.json()["documents"][0]["approval_status"] == "needs_review"

    assert all("oldcontext-token" not in excerpt for excerpt in _search_excerpts("oldcontext-token"))
    assert all("newcontext-token" not in excerpt for excerpt in _search_excerpts("newcontext-token"))


def test_embedding_failure_returns_readable_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    created = client.post(
        "/knowledge/ingestion-jobs",
        json={
            "source_name": "Embedding Failure Upload",
            "files": [
                {
                    "filename": "failure-sop.md",
                    "content": "# Failure SOP\nThis should fail embedding.",
                },
            ],
        },
    ).json()
    document_id = created["documents"][0]["id"]

    def fail_embed(value: str) -> list[float]:
        raise RuntimeError("OpenAI embedding request failed: ConnectError")

    monkeypatch.setattr(knowledge_ingestion_service.embedding_service, "embed_text", fail_embed)

    response = client.post(f"/knowledge/documents/{document_id}/approve")

    assert response.status_code == 502
    assert "OpenAI embedding request failed" in response.json()["detail"]


def test_extract_pdf_with_docling_mocked(monkeypatch) -> None:
    # docling is an optional dependency group (not installed in CI / not in the default
    # `uv sync`). This test mocks the docling converter, but monkeypatch.setattr still has to
    # import `docling.document_converter` to resolve the attribute path — so skip cleanly when
    # the optional extra is absent. The production path (_extract_pdf_with_docling) degrades
    # to a docling_unavailable ExtractionError without the package, unit-tested separately.
    pytest.importorskip("docling")
    from unittest.mock import MagicMock

    mock_converter = MagicMock()
    mock_result = MagicMock()
    mock_result.document.export_to_markdown.return_value = "# Mocked Docling Content\nDocling layout parsed successfully."
    mock_converter.convert.return_value = mock_result

    monkeypatch.setattr("docling.document_converter.DocumentConverter", lambda: mock_converter)

    from agent_studio.config import Settings
    settings = Settings(sagad_docling_enabled=True)

    res = extract_file(
        KnowledgeIngestionFile(
            filename="test_docling.pdf",
            content=_base64_bytes(b"PDF header dummy content"),
            encoding="base64"
        ),
        settings=settings
    )

    assert res.title == "Mocked Docling Content"
    assert "Docling layout parsed successfully" in res.content
    assert res.metadata["extractor"] == "pdf_docling"
