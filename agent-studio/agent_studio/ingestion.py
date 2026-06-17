from __future__ import annotations

import base64
import csv
from dataclasses import dataclass
from datetime import datetime, timezone
import io
import json
from pathlib import PurePosixPath
import re
from typing import Protocol

from psycopg.types.json import Jsonb

from agent_studio.config import Settings, get_settings
from agent_studio.db import (
    TrustedContext,
    connect,
    database_configured,
    initialize_database,
    resolve_trusted_context,
    set_app_context,
)
from agent_studio.embeddings import (
    EmbeddingService,
    content_hash,
    tokenize,
    vector_literal,
)
from agent_studio.knowledge import KnowledgeRecord
from agent_studio.schemas import (
    KnowledgeDocumentRecord,
    KnowledgeIngestionErrorRecord,
    KnowledgeIngestionFile,
    KnowledgeIngestionJobCreateRequest,
    KnowledgeIngestionJobRecord,
    KnowledgeIngestionJobResponse,
    KnowledgeSourceRecord,
)
from agent_studio.store import StoreContext


SUPPORTED_TEXT_EXTENSIONS = {".md", ".markdown", ".txt", ".log"}
SUPPORTED_TRANSCRIPT_EXTENSIONS = {".json", ".vtt", ".srt"}
SUPPORTED_STRUCTURED_EXTENSIONS = {".csv", ".xlsx"}
SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
DEFAULT_PACK_SLUG = "manual-ingestion"


class ExtractionError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


@dataclass(frozen=True)
class ExtractedDocument:
    title: str
    content: str
    metadata: dict[str, object]


class KnowledgeIngestionStoreProtocol(Protocol):
    backend_name: str

    def upsert_source(
        self,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeSourceRecord:
        ...

    def create_job(
        self,
        source: KnowledgeSourceRecord,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobRecord:
        ...

    def save_job(
        self,
        job: KnowledgeIngestionJobRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobRecord:
        ...

    def record_error(
        self,
        error: KnowledgeIngestionErrorRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionErrorRecord:
        ...

    def save_document(
        self,
        document: KnowledgeDocumentRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord:
        ...

    def list_sources(self, context: StoreContext | None = None) -> list[KnowledgeSourceRecord]:
        ...

    def touch_source(
        self,
        source_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeSourceRecord | None:
        ...

    def list_jobs(self, context: StoreContext | None = None) -> list[KnowledgeIngestionJobRecord]:
        ...

    def list_documents(
        self,
        source_id: str | None = None,
        context: StoreContext | None = None,
    ) -> list[KnowledgeDocumentRecord]:
        ...

    def get_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        ...

    def approve_document(
        self,
        document_id: str,
        embedding_service: EmbeddingService,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        ...

    def resync_document(
        self,
        document_id: str,
        embedding_service: EmbeddingService,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        ...

    def archive_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        ...

    def clear(self) -> None:
        ...


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _trusted_context(context: StoreContext | None) -> TrustedContext:
    scoped = context or StoreContext()
    return TrustedContext(
        organization_id=scoped.organization_id,
        user_id=scoped.user_id,
        role=scoped.role,
    )


def _stable_id(prefix: str, value: str, length: int = 16) -> str:
    return f"{prefix}_{content_hash(value)[:length]}"


def _clean_text(value: str) -> str:
    cleaned = value.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError(
        "decode_failed",
        "File text could not be decoded as UTF-8 or Windows-1252.",
    )


def _file_bytes(file: KnowledgeIngestionFile) -> bytes:
    if file.encoding == "base64":
        try:
            return base64.b64decode(file.content, validate=True)
        except ValueError as exc:
            raise ExtractionError(
                "invalid_base64",
                "File content was marked base64 but could not be decoded.",
            ) from exc
    return file.content.encode("utf-8")


def _title_from_content(content: str, fallback: str) -> str:
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped.removeprefix("# ").strip()
        if stripped:
            return stripped[:80]
    return fallback


def _category_from_filename(filename: str, explicit: str | None) -> str:
    if explicit:
        return explicit
    name = filename.lower()
    if "compliance" in name or "policy" in name:
        return "compliance"
    if "sop" in name or "process" in name or "procedure" in name:
        return "sops"
    if "qa" in name or "rubric" in name or "scorecard" in name:
        return "qa"
    if any(term in name for term in ("refund", "cancel", "angry", "escalation")):
        return "escalations"
    if "template" in name or "macro" in name:
        return "approved_templates"
    return "kb"


def _normalize_source_path(filename: str, source_path: str | None) -> str:
    candidate = source_path or filename
    return str(PurePosixPath(candidate.replace("\\", "/")))


def _pack_slug_for_source(source_id: str) -> str:
    return f"{DEFAULT_PACK_SLUG}:{source_id}"


def _extract_transcript_json(raw: bytes) -> ExtractedDocument:
    try:
        payload = json.loads(_decode_text(raw))
    except json.JSONDecodeError as exc:
        raise ExtractionError("invalid_transcript_json", "Transcript JSON is invalid.") from exc

    lines: list[str] = []
    if isinstance(payload, list):
        records = payload
    elif isinstance(payload, dict):
        records_value = payload.get("messages") or payload.get("transcript") or payload.get("items")
        records = records_value if isinstance(records_value, list) else [payload]
    else:
        records = []

    for item in records:
        if not isinstance(item, dict):
            continue
        speaker = item.get("speaker") or item.get("sender") or item.get("role") or "speaker"
        text = item.get("text") or item.get("content") or item.get("body") or ""
        timestamp = item.get("timestamp") or item.get("time") or item.get("created_at")
        prefix = f"[{timestamp}] {speaker}" if timestamp else str(speaker)
        if text:
            lines.append(f"{prefix}: {text}")

    content = _clean_text("\n".join(lines))
    if not content:
        raise ExtractionError("empty_transcript", "Transcript JSON did not contain readable messages.")
    return ExtractedDocument(
        title="Imported Transcript",
        content=content,
        metadata={"extractor": "transcript_json"},
    )


def _extract_vtt_or_srt(raw: bytes) -> ExtractedDocument:
    content = _decode_text(raw)
    lines: list[str] = []
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped or stripped.isdigit() or "-->" in stripped or stripped.upper() == "WEBVTT":
            continue
        lines.append(stripped)
    cleaned = _clean_text("\n".join(lines))
    if not cleaned:
        raise ExtractionError("empty_transcript", "Transcript file did not contain readable cue text.")
    return ExtractedDocument(
        title=_title_from_content(cleaned, "Imported Transcript"),
        content=cleaned,
        metadata={"extractor": "timed_transcript"},
    )


def _extract_csv(raw: bytes) -> ExtractedDocument:
    text = _decode_text(raw)
    reader = csv.DictReader(io.StringIO(text))
    rows: list[str] = []
    if reader.fieldnames:
        for index, row in enumerate(reader, start=1):
            cells = [f"{key}: {value}" for key, value in row.items() if value]
            if cells:
                rows.append(f"Row {index}: " + "; ".join(cells))
    else:
        simple_reader = csv.reader(io.StringIO(text))
        for index, row in enumerate(simple_reader, start=1):
            rows.append(f"Row {index}: " + "; ".join(cell for cell in row if cell))
    content = _clean_text("\n".join(rows))
    if not content:
        raise ExtractionError("empty_csv", "CSV file did not contain readable rows.")
    return ExtractedDocument(
        title="Imported Spreadsheet",
        content=content,
        metadata={"extractor": "csv"},
    )


def _extract_xlsx(raw: bytes) -> ExtractedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("missing_xlsx_parser", "XLSX parsing dependency is not installed.") from exc

    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(value).strip() if value is not None else "" for value in rows[0]]
        for row_index, row in enumerate(rows[1:], start=2):
            cells: list[str] = []
            for column_index, value in enumerate(row):
                if value is None:
                    continue
                header = headers[column_index] if column_index < len(headers) and headers[column_index] else f"Column {column_index + 1}"
                cells.append(f"{header}: {value}")
            if cells:
                lines.append(f"{sheet.title} row {row_index}: " + "; ".join(cells))
    content = _clean_text("\n".join(lines))
    if not content:
        raise ExtractionError("empty_xlsx", "XLSX file did not contain readable rows.")
    return ExtractedDocument(
        title="Imported Spreadsheet",
        content=content,
        metadata={"extractor": "xlsx", "sheets": [sheet.title for sheet in workbook.worksheets]},
    )


def _extract_docx(raw: bytes) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("missing_docx_parser", "DOCX parsing dependency is not installed.") from exc

    document = Document(io.BytesIO(raw))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    content = _clean_text("\n".join(lines))
    if not content:
        raise ExtractionError("empty_docx", "DOCX file did not contain readable text.")
    return ExtractedDocument(
        title=_title_from_content(content, "Imported Document"),
        content=content,
        metadata={"extractor": "docx"},
    )


def _fallback_pdf_text(raw: bytes) -> str:
    decoded = raw.decode("latin-1", errors="ignore")
    matches = re.findall(r"\(([^()]*)\)\s*Tj", decoded)
    return _clean_text("\n".join(match.replace("\\(", "(").replace("\\)", ")") for match in matches))


def _convert_pdf_to_images(raw: bytes, max_pages: int) -> list[object]:
    try:
        from pdf2image import convert_from_bytes
    except ImportError as exc:
        raise RuntimeError("pdf2image is not installed") from exc
    try:
        return list(convert_from_bytes(raw, first_page=1, last_page=max(1, max_pages)))
    except Exception as exc:
        raise RuntimeError(str(exc) or "PDF pages could not be rendered for OCR") from exc


def _image_to_text(image: object, lang: str, timeout_seconds: float) -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise RuntimeError("pytesseract is not installed") from exc
    try:
        return str(
            pytesseract.image_to_string(
                image,
                lang=lang,
                timeout=int(max(1, timeout_seconds)),
            ),
        )
    except Exception as exc:
        raise RuntimeError(str(exc) or "Tesseract OCR failed") from exc


def _extract_pdf_with_ocr(
    raw: bytes,
    settings: Settings,
    parser_error: str | None = None,
) -> ExtractedDocument:
    if not settings.sagad_ocr_enabled:
        detail = "PDF did not contain extractable text. Enable SAGAD_OCR_ENABLED to run local OCR."
        if parser_error:
            detail = f"{detail} PDF parser detail: {parser_error}"
        raise ExtractionError(
            "ocr_required",
            detail,
        )
    try:
        images = _convert_pdf_to_images(raw, settings.sagad_ocr_max_pages)
    except RuntimeError as exc:
        raise ExtractionError(
            "ocr_unavailable",
            f"OCR runtime is unavailable: {exc}",
        ) from exc
    text_parts: list[str] = []
    for image in images:
        try:
            page_text = _image_to_text(
                image,
                settings.sagad_ocr_lang,
                settings.sagad_ocr_timeout_seconds,
            )
        except RuntimeError as exc:
            raise ExtractionError(
                "ocr_failed",
                f"OCR failed while reading PDF page: {exc}",
            ) from exc
        if page_text.strip():
            text_parts.append(page_text)
    content = _clean_text("\n\n".join(text_parts))
    if not content:
        raise ExtractionError(
            "ocr_failed",
            "OCR completed but did not find readable text in the PDF.",
        )
    return ExtractedDocument(
        title=_title_from_content(content, "Imported OCR PDF"),
        content=content,
        metadata={
            "extractor": "pdf_ocr",
            "ocr_used": True,
            "ocr_pages": len(images),
            "ocr_lang": settings.sagad_ocr_lang,
            **({"pdf_parser_error": parser_error} if parser_error else {}),
        },
    )


def _extract_pdf_with_docling(
    raw: bytes,
    settings: Settings,
) -> ExtractedDocument:
    import tempfile
    import os
    from pathlib import Path

    # Create temporary scratch folder under workspace
    workspace_dir = Path(__file__).resolve().parents[2] / "scratch"
    workspace_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.NamedTemporaryFile(dir=str(workspace_dir), suffix=".pdf", delete=False) as tmp:
        tmp.write(raw)
        tmp_path = tmp.name

    try:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(tmp_path)

        try:
            content = result.document.export_to_markdown()
        except AttributeError:
            try:
                content = result.render_as_markdown()
            except AttributeError:
                content = str(result.document)

        cleaned_content = _clean_text(content)

        return ExtractedDocument(
            title=_title_from_content(cleaned_content, "Imported Docling PDF"),
            content=cleaned_content,
            metadata={
                "extractor": "pdf_docling",
                "ocr_used": True,
            },
        )
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _extract_pdf(raw: bytes, settings: Settings | None = None) -> ExtractedDocument:
    actual_settings = settings or get_settings()
    parser_error: str | None = None

    if actual_settings.sagad_docling_enabled:
        try:
            return _extract_pdf_with_docling(raw, actual_settings)
        except Exception as exc:
            parser_error = f"Docling failed: {exc}"

    text_parts: list[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    except Exception as exc:
        parser_error = str(exc) or exc.__class__.__name__
        text_parts = []

    content = _clean_text("\n".join(text_parts)) or _fallback_pdf_text(raw)
    if not content:
        return _extract_pdf_with_ocr(raw, actual_settings, parser_error=parser_error)
    return ExtractedDocument(
        title=_title_from_content(content, "Imported PDF"),
        content=content,
        metadata={"extractor": "pdf", "ocr_used": False, **({"pdf_parser_error": parser_error} if parser_error else {})},
    )


def extract_file(
    file: KnowledgeIngestionFile,
    settings: Settings | None = None,
) -> ExtractedDocument:
    raw = _file_bytes(file)
    extension = PurePosixPath(file.filename.lower()).suffix
    if extension in SUPPORTED_TEXT_EXTENSIONS:
        content = _clean_text(_decode_text(raw))
        if not content:
            raise ExtractionError("empty_text", "Text file did not contain readable content.")
        return ExtractedDocument(
            title=_title_from_content(content, PurePosixPath(file.filename).stem),
            content=content,
            metadata={"extractor": "text"},
        )
    if extension == ".json":
        return _extract_transcript_json(raw)
    if extension in {".vtt", ".srt"}:
        return _extract_vtt_or_srt(raw)
    if extension == ".csv":
        return _extract_csv(raw)
    if extension == ".xlsx":
        return _extract_xlsx(raw)
    if extension == ".docx":
        return _extract_docx(raw)
    if extension == ".pdf":
        return _extract_pdf(raw, settings=settings)
    raise ExtractionError(
        "unsupported_file_type",
        f"Unsupported knowledge file type '{extension or 'unknown'}'.",
    )


def build_chunks(document: KnowledgeDocumentRecord, max_tokens: int = 700) -> list[tuple[str, str, int]]:
    paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", document.content) if paragraph.strip()]
    chunks: list[tuple[str, str, int]] = []
    current: list[str] = []
    current_tokens = 0
    for paragraph in paragraphs or [document.content]:
        token_count = len(tokenize(paragraph))
        if current and current_tokens + token_count > max_tokens:
            content = "\n\n".join(current)
            chunks.append((document.title if not chunks else None or document.title, content, current_tokens))
            current = []
            current_tokens = 0
        current.append(paragraph)
        current_tokens += token_count
    if current:
        content = "\n\n".join(current)
        chunks.append((document.title, content, current_tokens))
    return chunks


class InMemoryKnowledgeIngestionStore:
    backend_name = "memory"

    def __init__(self) -> None:
        self._sources: dict[str, KnowledgeSourceRecord] = {}
        self._jobs: dict[str, KnowledgeIngestionJobRecord] = {}
        self._documents: dict[str, KnowledgeDocumentRecord] = {}
        self._errors: list[KnowledgeIngestionErrorRecord] = []

    def upsert_source(
        self,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeSourceRecord:
        source_id = _stable_id("ksrc", f"{request.source_type}:{request.source_name}")
        now = _now()
        existing = self._sources.get(source_id)
        source = KnowledgeSourceRecord(
            id=source_id,
            source_type=request.source_type,
            name=request.source_name,
            status="ready",
            sync_policy="manual",
            last_synced_at=now,
            metadata=request.metadata,
            created_at=existing.created_at if existing else now,
            updated_at=now,
        )
        self._sources[source_id] = source
        return source

    def create_job(
        self,
        source: KnowledgeSourceRecord,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobRecord:
        job = KnowledgeIngestionJobRecord(
            source_id=source.id,
            source_name=source.name,
            source_type=source.source_type,
            status="queued",
            total_files=len(request.files),
            metadata=request.metadata,
        )
        self._jobs[job.id] = job
        return job

    def save_job(
        self,
        job: KnowledgeIngestionJobRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobRecord:
        job.updated_at = _now()
        job.errors = [error for error in self._errors if error.job_id == job.id]
        self._jobs[job.id] = job
        return job

    def record_error(
        self,
        error: KnowledgeIngestionErrorRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionErrorRecord:
        self._errors.append(error)
        return error

    def save_document(
        self,
        document: KnowledgeDocumentRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord:
        existing = self._documents.get(document.id)
        if existing and existing.content_hash != document.content_hash:
            document.version = existing.version + 1
        elif existing:
            document.version = existing.version
            document.created_at = existing.created_at
        document.updated_at = _now()
        self._documents[document.id] = document
        return document

    def list_sources(self, context: StoreContext | None = None) -> list[KnowledgeSourceRecord]:
        return sorted(self._sources.values(), key=lambda source: source.updated_at, reverse=True)

    def touch_source(
        self,
        source_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeSourceRecord | None:
        source = self._sources.get(source_id)
        if source is None:
            return None
        source.last_synced_at = _now()
        source.updated_at = source.last_synced_at
        self._sources[source_id] = source
        return source

    def list_jobs(self, context: StoreContext | None = None) -> list[KnowledgeIngestionJobRecord]:
        return sorted(self._jobs.values(), key=lambda job: job.created_at, reverse=True)

    def list_documents(
        self,
        source_id: str | None = None,
        context: StoreContext | None = None,
    ) -> list[KnowledgeDocumentRecord]:
        documents = list(self._documents.values())
        if source_id is not None:
            documents = [document for document in documents if document.source_id == source_id]
        return sorted(documents, key=lambda document: document.updated_at, reverse=True)

    def get_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        return self._documents.get(document_id)

    def approve_document(
        self,
        document_id: str,
        embedding_service: EmbeddingService,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        document = self._documents.get(document_id)
        if document is None:
            return None
        chunks = build_chunks(document)
        for _, chunk_content, _ in chunks:
            embedding_service.embed_text(chunk_content)
        document.approval_status = "approved"
        document.chunk_count = len(chunks)
        document.updated_at = _now()
        self._documents[document_id] = document
        return document

    def resync_document(
        self,
        document_id: str,
        embedding_service: EmbeddingService,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        document = self._documents.get(document_id)
        if document is None:
            return None
        document.metadata = {**document.metadata, "last_resynced_at": _now().isoformat()}
        document.updated_at = _now()
        if document.approval_status == "approved":
            chunks = build_chunks(document)
            for _, chunk_content, _ in chunks:
                embedding_service.embed_text(chunk_content)
            document.chunk_count = len(chunks)
        self._documents[document_id] = document
        return document

    def archive_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        document = self._documents.get(document_id)
        if document is None:
            return None
        document.approval_status = "archived"
        document.chunk_count = 0
        document.updated_at = _now()
        self._documents[document_id] = document
        return document

    def clear(self) -> None:
        self._sources.clear()
        self._jobs.clear()
        self._documents.clear()
        self._errors.clear()


class PostgresKnowledgeIngestionStore:
    backend_name = "postgres"

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        initialize_database(settings)

    def upsert_source(
        self,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeSourceRecord:
        source_id = _stable_id("ksrc", f"{request.source_type}:{request.source_name}")
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                INSERT INTO knowledge_sources (
                  id,
                  organization_id,
                  source_type,
                  name,
                  status,
                  sync_policy,
                  last_synced_at,
                  metadata
                )
                VALUES (%s, %s, %s, %s, 'ready', 'manual', now(), %s)
                ON CONFLICT (organization_id, source_type, name) DO UPDATE SET
                  status = EXCLUDED.status,
                  last_synced_at = EXCLUDED.last_synced_at,
                  metadata = EXCLUDED.metadata,
                  updated_at = now()
                RETURNING *
                """,
                (
                    source_id,
                    scoped.organization_id,
                    request.source_type,
                    request.source_name,
                    Jsonb(request.metadata),
                ),
            ).fetchone()
            connection.commit()
        return _source_from_row(row)

    def create_job(
        self,
        source: KnowledgeSourceRecord,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobRecord:
        job = KnowledgeIngestionJobRecord(
            source_id=source.id,
            source_name=source.name,
            source_type=source.source_type,
            total_files=len(request.files),
            metadata=request.metadata,
        )
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                INSERT INTO knowledge_ingestion_jobs (
                  id,
                  organization_id,
                  source_id,
                  source_name,
                  source_type,
                  status,
                  total_files,
                  processed_files,
                  failed_files,
                  summary,
                  metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING *
                """,
                (
                    job.id,
                    scoped.organization_id,
                    source.id,
                    source.name,
                    source.source_type,
                    job.status,
                    job.total_files,
                    job.processed_files,
                    job.failed_files,
                    job.summary,
                    Jsonb(job.metadata),
                ),
            ).fetchone()
            connection.commit()
        return _job_from_row(row, [])

    def save_job(
        self,
        job: KnowledgeIngestionJobRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobRecord:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                UPDATE knowledge_ingestion_jobs
                SET status = %s,
                    processed_files = %s,
                    failed_files = %s,
                    summary = %s,
                    metadata = %s,
                    updated_at = now()
                WHERE organization_id = %s
                  AND id = %s
                RETURNING *
                """,
                (
                    job.status,
                    job.processed_files,
                    job.failed_files,
                    job.summary,
                    Jsonb(job.metadata),
                    scoped.organization_id,
                    job.id,
                ),
            ).fetchone()
            errors = self._errors_for_job(connection, scoped.organization_id, job.id)
            connection.commit()
        return _job_from_row(row, errors) if row else job

    def record_error(
        self,
        error: KnowledgeIngestionErrorRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionErrorRecord:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                INSERT INTO knowledge_ingestion_errors (
                  id,
                  organization_id,
                  job_id,
                  source_path,
                  error_code,
                  message,
                  metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                RETURNING *
                """,
                (
                    error.id,
                    scoped.organization_id,
                    error.job_id,
                    error.source_path,
                    error.error_code,
                    error.message,
                    Jsonb(error.metadata),
                ),
            ).fetchone()
            connection.commit()
        return _error_from_row(row)

    def save_document(
        self,
        document: KnowledgeDocumentRecord,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            existing = connection.execute(
                """
                SELECT *
                FROM knowledge_documents
                WHERE organization_id = %s
                  AND pack_slug = %s
                  AND source_path = %s
                """,
                (scoped.organization_id, document.pack_slug, document.source_path),
            ).fetchone()
            version = 1
            if existing is not None:
                existing_version = existing.get("version")
                version = int(existing_version) if isinstance(existing_version, int) else 1
                if str(existing["content_hash"]) != document.content_hash:
                    version += 1
            row = connection.execute(
                """
                INSERT INTO knowledge_documents (
                  id,
                  organization_id,
                  source_id,
                  last_ingestion_job_id,
                  pack_slug,
                  category,
                  source_path,
                  title,
                  content,
                  content_hash,
                  version,
                  approval_status,
                  metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'needs_review', %s)
                ON CONFLICT (organization_id, pack_slug, source_path) DO UPDATE SET
                  source_id = EXCLUDED.source_id,
                  last_ingestion_job_id = EXCLUDED.last_ingestion_job_id,
                  category = EXCLUDED.category,
                  title = EXCLUDED.title,
                  content = EXCLUDED.content,
                  content_hash = EXCLUDED.content_hash,
                  version = EXCLUDED.version,
                  approval_status = 'needs_review',
                  metadata = EXCLUDED.metadata,
                  updated_at = now()
                RETURNING *
                """,
                (
                    document.id if existing is None else existing["id"],
                    scoped.organization_id,
                    document.source_id,
                    document.job_id,
                    document.pack_slug,
                    document.category,
                    document.source_path,
                    document.title,
                    document.content,
                    document.content_hash,
                    version,
                    Jsonb(document.metadata),
                ),
            ).fetchone()
            connection.commit()
        return _document_from_row(row, chunk_count=0)

    def list_sources(self, context: StoreContext | None = None) -> list[KnowledgeSourceRecord]:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            rows = connection.execute(
                """
                SELECT *
                FROM knowledge_sources
                WHERE organization_id = %s
                ORDER BY updated_at DESC
                """,
                (scoped.organization_id,),
            ).fetchall()
        return [_source_from_row(row) for row in rows]

    def touch_source(
        self,
        source_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeSourceRecord | None:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                UPDATE knowledge_sources
                SET last_synced_at = now(),
                    updated_at = now()
                WHERE organization_id = %s
                  AND id = %s
                RETURNING *
                """,
                (scoped.organization_id, source_id),
            ).fetchone()
            connection.commit()
        return _source_from_row(row) if row else None

    def list_jobs(self, context: StoreContext | None = None) -> list[KnowledgeIngestionJobRecord]:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            rows = connection.execute(
                """
                SELECT *
                FROM knowledge_ingestion_jobs
                WHERE organization_id = %s
                ORDER BY created_at DESC
                """,
                (scoped.organization_id,),
            ).fetchall()
            errors = {
                job_id: self._errors_for_job(connection, scoped.organization_id, job_id)
                for job_id in [str(row["id"]) for row in rows]
            }
        return [_job_from_row(row, errors.get(str(row["id"]), [])) for row in rows]

    def list_documents(
        self,
        source_id: str | None = None,
        context: StoreContext | None = None,
    ) -> list[KnowledgeDocumentRecord]:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            source_filter = "AND source_id = %s" if source_id is not None else ""
            params: tuple[object, ...] = (
                (scoped.organization_id, source_id)
                if source_id is not None
                else (scoped.organization_id,)
            )
            rows = connection.execute(
                f"""
                SELECT
                  knowledge_documents.*,
                  (
                    SELECT count(*)
                    FROM knowledge_chunks
                    WHERE knowledge_chunks.document_id = knowledge_documents.id
                  ) AS chunk_count
                FROM knowledge_documents
                WHERE organization_id = %s
                  {source_filter}
                ORDER BY updated_at DESC
                """,
                params,
            ).fetchall()
        return [_document_from_row(row, chunk_count=int(row["chunk_count"])) for row in rows]

    def get_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                SELECT
                  knowledge_documents.*,
                  (
                    SELECT count(*)
                    FROM knowledge_chunks
                    WHERE knowledge_chunks.document_id = knowledge_documents.id
                  ) AS chunk_count
                FROM knowledge_documents
                WHERE organization_id = %s
                  AND id = %s
                """,
                (scoped.organization_id, document_id),
            ).fetchone()
        return _document_from_row(row, chunk_count=int(row["chunk_count"])) if row else None

    def approve_document(
        self,
        document_id: str,
        embedding_service: EmbeddingService,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                SELECT *
                FROM knowledge_documents
                WHERE organization_id = %s
                  AND id = %s
                """,
                (scoped.organization_id, document_id),
            ).fetchone()
            if row is None:
                return None
            document = _document_from_row(row, chunk_count=0)
            connection.execute(
                "DELETE FROM knowledge_chunks WHERE organization_id = %s AND document_id = %s",
                (scoped.organization_id, document.id),
            )
            chunks = build_chunks(document)
            for index, (heading, chunk_content, token_count) in enumerate(chunks):
                chunk_id = f"{document.id}:chunk:{index}"
                chunk_hash = content_hash(chunk_content)
                connection.execute(
                    """
                    INSERT INTO knowledge_chunks (
                      id,
                      organization_id,
                      document_id,
                      chunk_index,
                      heading,
                      content,
                      content_hash,
                      token_count,
                      metadata
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """,
                    (
                        chunk_id,
                        scoped.organization_id,
                        document.id,
                        index,
                        heading,
                        chunk_content,
                        chunk_hash,
                        token_count,
                        Jsonb({"source_path": document.source_path, "version": document.version}),
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO knowledge_chunk_embeddings (
                      chunk_id,
                      organization_id,
                      embedding_model,
                      embedding,
                      content_hash
                    )
                    VALUES (%s, %s, %s, %s::vector, %s)
                    ON CONFLICT (chunk_id, embedding_model) DO UPDATE SET
                      embedding = EXCLUDED.embedding,
                      content_hash = EXCLUDED.content_hash,
                      updated_at = now()
                    """,
                    (
                        chunk_id,
                        scoped.organization_id,
                        embedding_service.embedding_model,
                        vector_literal(embedding_service.embed_text(chunk_content)),
                        chunk_hash,
                    ),
                )
            updated = connection.execute(
                """
                UPDATE knowledge_documents
                SET approval_status = 'approved',
                    updated_at = now()
                WHERE organization_id = %s
                  AND id = %s
                RETURNING *
                """,
                (scoped.organization_id, document.id),
            ).fetchone()
            connection.commit()
        return _document_from_row(updated, chunk_count=len(chunks)) if updated else None

    def resync_document(
        self,
        document_id: str,
        embedding_service: EmbeddingService,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            row = connection.execute(
                """
                SELECT *
                FROM knowledge_documents
                WHERE organization_id = %s
                  AND id = %s
                """,
                (scoped.organization_id, document_id),
            ).fetchone()
            if row is None:
                return None
            document = _document_from_row(row, chunk_count=0)
            resync_metadata = {
                **document.metadata,
                "last_resynced_at": _now().isoformat(),
            }
            chunk_count = 0
            if document.approval_status == "approved":
                connection.execute(
                    "DELETE FROM knowledge_chunks WHERE organization_id = %s AND document_id = %s",
                    (scoped.organization_id, document.id),
                )
                chunks = build_chunks(document)
                chunk_count = len(chunks)
                for index, (heading, chunk_content, token_count) in enumerate(chunks):
                    chunk_id = f"{document.id}:chunk:{index}"
                    chunk_hash = content_hash(chunk_content)
                    connection.execute(
                        """
                        INSERT INTO knowledge_chunks (
                          id,
                          organization_id,
                          document_id,
                          chunk_index,
                          heading,
                          content,
                          content_hash,
                          token_count,
                          metadata
                        )
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """,
                        (
                            chunk_id,
                            scoped.organization_id,
                            document.id,
                            index,
                            heading,
                            chunk_content,
                            chunk_hash,
                            token_count,
                            Jsonb({"source_path": document.source_path, "version": document.version}),
                        ),
                    )
                    connection.execute(
                        """
                        INSERT INTO knowledge_chunk_embeddings (
                          chunk_id,
                          organization_id,
                          embedding_model,
                          embedding,
                          content_hash
                        )
                        VALUES (%s, %s, %s, %s::vector, %s)
                        ON CONFLICT (chunk_id, embedding_model) DO UPDATE SET
                          embedding = EXCLUDED.embedding,
                          content_hash = EXCLUDED.content_hash,
                          updated_at = now()
                        """,
                        (
                            chunk_id,
                            scoped.organization_id,
                            embedding_service.embedding_model,
                            vector_literal(embedding_service.embed_text(chunk_content)),
                            chunk_hash,
                        ),
                    )
            updated = connection.execute(
                """
                UPDATE knowledge_documents
                SET metadata = %s,
                    updated_at = now()
                WHERE organization_id = %s
                  AND id = %s
                RETURNING *
                """,
                (Jsonb(resync_metadata), scoped.organization_id, document.id),
            ).fetchone()
            connection.commit()
        return _document_from_row(updated, chunk_count=chunk_count) if updated else None

    def archive_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        with connect(self.settings) as connection:
            scoped = resolve_trusted_context(connection, _trusted_context(context))
            set_app_context(connection, scoped)
            connection.execute(
                "DELETE FROM knowledge_chunks WHERE organization_id = %s AND document_id = %s",
                (scoped.organization_id, document_id),
            )
            row = connection.execute(
                """
                UPDATE knowledge_documents
                SET approval_status = 'archived',
                    updated_at = now()
                WHERE organization_id = %s
                  AND id = %s
                RETURNING *
                """,
                (scoped.organization_id, document_id),
            ).fetchone()
            connection.commit()
        return _document_from_row(row, chunk_count=0) if row else None

    def clear(self) -> None:
        with connect(self.settings) as connection:
            connection.execute("TRUNCATE knowledge_ingestion_errors CASCADE")
            connection.execute("TRUNCATE knowledge_ingestion_jobs CASCADE")
            connection.execute("TRUNCATE knowledge_sources CASCADE")
            connection.commit()

    def _errors_for_job(
        self,
        connection: object,
        organization_id: str | None,
        job_id: str,
    ) -> list[KnowledgeIngestionErrorRecord]:
        rows = connection.execute(
            """
            SELECT *
            FROM knowledge_ingestion_errors
            WHERE organization_id = %s
              AND job_id = %s
            ORDER BY created_at ASC
            """,
            (organization_id, job_id),
        ).fetchall()
        return [_error_from_row(row) for row in rows]


def _source_from_row(row: object) -> KnowledgeSourceRecord:
    payload = dict(row)  # type: ignore[arg-type]
    return KnowledgeSourceRecord(
        id=str(payload["id"]),
        source_type=payload["source_type"],
        name=str(payload["name"]),
        status=payload["status"],
        sync_policy=str(payload["sync_policy"]),
        last_synced_at=payload["last_synced_at"],
        metadata=payload["metadata"] if isinstance(payload["metadata"], dict) else {},
        created_at=payload["created_at"],
        updated_at=payload["updated_at"],
    )


def _job_from_row(
    row: object,
    errors: list[KnowledgeIngestionErrorRecord],
) -> KnowledgeIngestionJobRecord:
    payload = dict(row)  # type: ignore[arg-type]
    return KnowledgeIngestionJobRecord(
        id=str(payload["id"]),
        source_id=str(payload["source_id"]),
        source_name=str(payload["source_name"]),
        source_type=payload["source_type"],
        status=payload["status"],
        total_files=int(payload["total_files"]),
        processed_files=int(payload["processed_files"]),
        failed_files=int(payload["failed_files"]),
        summary=str(payload["summary"]),
        metadata=payload["metadata"] if isinstance(payload["metadata"], dict) else {},
        errors=errors,
        created_at=payload["created_at"],
        updated_at=payload["updated_at"],
    )


def _error_from_row(row: object) -> KnowledgeIngestionErrorRecord:
    payload = dict(row)  # type: ignore[arg-type]
    return KnowledgeIngestionErrorRecord(
        id=str(payload["id"]),
        job_id=str(payload["job_id"]),
        source_path=str(payload["source_path"]),
        error_code=str(payload["error_code"]),
        message=str(payload["message"]),
        metadata=payload["metadata"] if isinstance(payload["metadata"], dict) else {},
        created_at=payload["created_at"],
    )


def _document_from_row(row: object, *, chunk_count: int) -> KnowledgeDocumentRecord:
    payload = dict(row)  # type: ignore[arg-type]
    return KnowledgeDocumentRecord(
        id=str(payload["id"]),
        source_id=str(payload["source_id"]) if payload.get("source_id") else None,
        job_id=str(payload["last_ingestion_job_id"]) if payload.get("last_ingestion_job_id") else None,
        pack_slug=str(payload["pack_slug"]),
        category=str(payload["category"]),
        source_path=str(payload["source_path"]),
        title=str(payload["title"]),
        content=str(payload["content"]),
        content_hash=str(payload["content_hash"]),
        version=int(payload["version"]),
        approval_status=payload["approval_status"],
        chunk_count=chunk_count,
        metadata=payload["metadata"] if isinstance(payload["metadata"], dict) else {},
        created_at=payload["created_at"],
        updated_at=payload["updated_at"],
    )


def _record_from_document(document: KnowledgeDocumentRecord) -> KnowledgeRecord:
    return KnowledgeRecord(
        id=document.id,
        title=document.title,
        category=document.category,
        source_path=document.source_path,
        content=document.content,
        approval_status=document.approval_status,
    )


class KnowledgeIngestionService:
    def __init__(
        self,
        store: KnowledgeIngestionStoreProtocol,
        settings: Settings,
        runtime_retriever: object | None = None,
    ) -> None:
        self.store = store
        self.settings = settings
        self.embedding_service = EmbeddingService(settings)
        self.runtime_retriever = runtime_retriever

    def ingest(
        self,
        request: KnowledgeIngestionJobCreateRequest,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobResponse:
        source = self.store.upsert_source(request, context=context)
        job = self.store.create_job(source, request, context=context)
        job.status = "extracting"
        self.store.save_job(job, context=context)

        documents: list[KnowledgeDocumentRecord] = []
        errors: list[KnowledgeIngestionErrorRecord] = []
        for file in request.files:
            source_path = _normalize_source_path(file.filename, file.source_path)
            try:
                extracted = extract_file(file, settings=self.settings)
                category = _category_from_filename(file.filename, file.category)
                content = _clean_text(extracted.content)
                document_hash = content_hash(content)
                document = KnowledgeDocumentRecord(
                    id=_stable_id("kdoc", f"{source.id}:{source_path}"),
                    source_id=source.id,
                    job_id=job.id,
                    pack_slug=_pack_slug_for_source(source.id),
                    category=category,
                    source_path=source_path,
                    title=extracted.title,
                    content=content,
                    content_hash=document_hash,
                    approval_status="needs_review",
                    metadata={
                        **file.metadata,
                        **extracted.metadata,
                        "original_filename": file.filename,
                        "requires_review": True,
                    },
                )
                self._remove_runtime_record(document.id)
                documents.append(self.store.save_document(document, context=context))
                job.processed_files += 1
            except ExtractionError as exc:
                error = KnowledgeIngestionErrorRecord(
                    job_id=job.id,
                    source_path=source_path,
                    error_code=exc.code,
                    message=exc.message,
                    metadata={"filename": file.filename},
                )
                errors.append(self.store.record_error(error, context=context))
                job.failed_files += 1

        job.status = "needs_review" if documents else "failed"
        job.summary = (
            f"{len(documents)} document(s) need review; {len(errors)} file(s) failed."
            if documents
            else f"No documents were ingested; {len(errors)} file(s) failed."
        )
        job = self.store.save_job(job, context=context)
        return KnowledgeIngestionJobResponse(job=job, documents=documents, errors=errors)

    def list_sources(
        self,
        context: StoreContext | None = None,
    ) -> list[KnowledgeSourceRecord]:
        return self.store.list_sources(context=context)

    def approve_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        document = self.store.approve_document(
            document_id,
            self.embedding_service,
            context=context,
        )
        if document is not None:
            self._register_runtime_record(document)
        return document

    def resync_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        document = self.store.resync_document(
            document_id,
            self.embedding_service,
            context=context,
        )
        if document is not None and document.approval_status == "approved":
            self._register_runtime_record(document)
        return document

    def sync_source(
        self,
        source_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeIngestionJobResponse | None:
        source = self.store.touch_source(source_id, context=context)
        if source is None:
            return None
        documents = self.store.list_documents(source_id=source.id, context=context)
        request = KnowledgeIngestionJobCreateRequest(
            source_name=source.name,
            source_type=source.source_type,
            files=[],
            metadata={"sync_kind": "local_reindex", "source_id": source.id},
        )
        job = self.store.create_job(source, request, context=context)
        job.status = "embedding"
        job.total_files = len(documents)
        synced_documents: list[KnowledgeDocumentRecord] = []
        errors: list[KnowledgeIngestionErrorRecord] = []
        for document in documents:
            try:
                synced = self.resync_document(document.id, context=context)
                if synced is not None:
                    synced_documents.append(synced)
                    job.processed_files += 1
            except RuntimeError as exc:
                error = KnowledgeIngestionErrorRecord(
                    job_id=job.id,
                    source_path=document.source_path,
                    error_code="resync_failed",
                    message=str(exc),
                    metadata={"document_id": document.id},
                )
                errors.append(self.store.record_error(error, context=context))
                job.failed_files += 1
        job.status = "failed" if errors and not synced_documents else "ready"
        if any(document.approval_status == "needs_review" for document in synced_documents):
            job.status = "needs_review"
        job.summary = (
            f"Local re-index refreshed {len(synced_documents)} document(s); {len(errors)} failed."
        )
        job = self.store.save_job(job, context=context)
        return KnowledgeIngestionJobResponse(
            job=job,
            documents=synced_documents,
            errors=errors,
        )

    def archive_document(
        self,
        document_id: str,
        context: StoreContext | None = None,
    ) -> KnowledgeDocumentRecord | None:
        document = self.store.archive_document(document_id, context=context)
        if document is not None:
            self._remove_runtime_record(document.id)
        return document

    def _register_runtime_record(self, document: KnowledgeDocumentRecord) -> None:
        if self.runtime_retriever is None:
            return
        add_record = getattr(self.runtime_retriever, "add_record", None)
        if callable(add_record):
            add_record(_record_from_document(document))
        fallback = getattr(self.runtime_retriever, "fallback", None)
        fallback_add_record = getattr(fallback, "add_record", None)
        if callable(fallback_add_record):
            fallback_add_record(_record_from_document(document))

    def _remove_runtime_record(self, document_id: str) -> None:
        if self.runtime_retriever is None:
            return
        remove_record = getattr(self.runtime_retriever, "remove_record", None)
        if callable(remove_record):
            remove_record(document_id)
        fallback = getattr(self.runtime_retriever, "fallback", None)
        fallback_remove_record = getattr(fallback, "remove_record", None)
        if callable(fallback_remove_record):
            fallback_remove_record(document_id)


def build_knowledge_ingestion_store(settings: Settings | None = None) -> KnowledgeIngestionStoreProtocol:
    scoped_settings = settings or get_settings()
    if database_configured(scoped_settings):
        return PostgresKnowledgeIngestionStore(scoped_settings)
    return InMemoryKnowledgeIngestionStore()
